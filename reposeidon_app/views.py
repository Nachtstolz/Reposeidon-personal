from turtle import title
from django.shortcuts import render, redirect
from .forms import UserDataForm
from django.utils import timezone #시간 설정 패키지
from django.http import HttpResponseRedirect
from django.http import JsonResponse
from .models import * #모든 모델 상속
from django.views import View
from datetime import datetime

#rest-framework import
from rest_framework import generics
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from .serializers import UserDataSerializer
from .serializers import FileSerializer
from rest_framework import viewsets

from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
import os
from rest_framework.viewsets import ViewSet
# from .serializers import UploadSerializer
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


import io
from rest_framework.parsers import JSONParser
from .krwdrnk_md import *



# Create your views here.

####html연결
# #메인페이지 연결
# def mainPage(request):
#     return render(request, 'reposeidon_app/base.html')

# #AJAX활용 비동기식
# def showOutput(request):
#     userContent = request.GET.get('userContent') #본문
#     userTitle = request.GET.get('userTitle') #제목
#     userWriter = request.GET.get('userWriter') #유저 이름

#     # 유저가 입력한 내용
#     thisUserData = UserData()
#     thisUserData.title = userTitle
#     thisUserData.content = userContent
#     thisUserData.writer = userWriter
#     thisUserData.created = timezone.now()
#     thisUserData.save() #유저의 입력내용 DB저장할 경우에 사용

#     # 크롤링 및 AI 모델 연결해서 데이터 처리 코드 입력


#     #결과 값을 배열형태로 output에 담아서 JSON형태로 반환
#     # Ajax 테스트용
#     output = f"로컬 저장완료 \n작성자: {thisUserData.writer} \n제목: {thisUserData.title} \n내용: {thisUserData.content} \n검색일시: {thisUserData.created.strftime('%Y-%m-%dT%H:%M:%S')}"

#     context = {
#         'output': output,
#         'outputCreated': thisUserData.created,
#     }

#     return JsonResponse(context)

def front(request):
    context = {}
    return render(request, "index.html", context)

target = ''

@api_view(['GET', 'POST'])
def listUserData(request):
    if request.method == 'POST':
        serializer = UserDataSerializer(data=request.data)

        if serializer.is_valid():
            #serializer.save()
        
            target = serializer.validated_data['title']
            content = serializer.validated_data['content']
            
            #AI모델 연결
            result = krwdrnk(content)
            print('result')
            print(result)

            serializer.validated_data['content'] = result
            print('serializer')
            print(serializer)
            serializer.save()

            return Response(status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'GET':
        allData = UserData.objects.all()
        # thisData = UserData.objects.filter(writer=target) 
        serializer = UserDataSerializer(allData, many=True)
        print("get")
        print(serializer)
        # thisData = allData.get(title=target)
        # serializer = UserDataSerializer(thisData, many=True)
        
        return Response(serializer.data)



@api_view(['DELETE'])
def userDataDetail(request, pk):
    try:
        thisData = UserData.objects.get(pk=pk)
    except UserData.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'DELETE':
        thisData.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


#파일 업로드 받기
# class FileUploadViewSet(viewsets.ViewSet):
#     def create(self, request):
#         serializer_class = FileSerializer(data=request.data)
#         if 'file' not in request.FILES or not serializer_class.is_valid():
#             return Response(status=status.HTTP_400_BAD_REQUEST)
#         else:
#             handle_uploaded_file(request.FILES['file'])
#             return Response(status=status.HTTP_201_CREATED)

# def handle_uploaded_file(f):
#     with open(f.name, 'wb+') as destination:
#         for chunk in f.chunks():
#             destination.write(chunk)

#파일 업로드 로직
class FileView(APIView):
  parser_classes = (MultiPartParser, FormParser)
  
  def post(self, request, *args, **kwargs):
    file_serializer = FileSerializer(data=request.data)
    print("file_serializer:", file_serializer)
    # print("file_serializer.data:", file_serializer.data)

    if file_serializer.is_valid():
        
        #기본적으로 media경로에 저장된다.
        file_serializer.save()
        print("file_serializer.data:", file_serializer.data)

        file_name = str(file_serializer.data['file'])
        file_path = os.getcwd()+'\\media\\'+file_name[1:]
        file_format = os.path.splitext(file_name)[1]

        print("file Path:",file_path )
        print("file Format:",file_format)
        
        if file_format == '.txt':
            with open(file_path, "r", encoding="utf-8" ) as file_data:
                content = str(file_data.read())

                #AI모델 연결
                AI_result = krwdrnk(content)
                print('AI_result')
                print(AI_result)

                #크롤링 모델에 넣고 결과 받기
                crawling_result = run_crawling(AI_result)

        if file_format == '.docx':
            file_data = Document(file_path)

            content = ''
            print(file_data)
            for x, paragraph in enumerate(file_data.paragraphs):
                print(str(x) + " : " + paragraph.text)
                content += str(paragraph.text)

            #AI모델 연결
            AI_result = krwdrnk(content)
            print('AI_result')
            print(AI_result)
            
            #크롤링 모델에 넣고 결과 받기
            crawling_result = run_crawling(AI_result)


        return Response(file_serializer.data, status=status.HTTP_201_CREATED)
    else:
        return Response(file_serializer.errors, status=status.HTTP_400_BAD_REQUEST)


def run_crawling(content):
    #크롤링 모델에 result넣기

    #크롤링 결과를 front로 전송
    crawling_result = ''
    

    return crawling_result




def return_file():
    doc = Document('1')
    tables = doc.tables

    # 1번째 테이블 0, 0 셀 - 2번째 paragraphs에 입력
    tables[0].cell(0, 0).paragraphs[1].text = 'paragraphs[1]'

    # 2번째 테이블 0, 0 셀에 입력
    tables[1].cell(0, 0).text = 'table[1] (0,0)'




# # ViewSets define the view behavior.
# class UploadViewSet(ViewSet):
#     serializer_class = UploadSerializer

#     def list(self, request):
#         return Response("GET API")

#     def create(self, request):
#         file_uploaded = request.FILES.post('file_uploaded')
#         print("file_uploaded:",file_uploaded)
#         content_type = file_uploaded.content_type
#         response = "POST API and you have uploaded a {} file".format(content_type)
#         return Response(response)